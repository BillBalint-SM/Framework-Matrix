from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(__file__).resolve().parents[2]
REFERENCE = WORKSPACE / "work" / "document-template" / "reference.docx"
OUTPUT = WORKSPACE / "outputs" / "sdd-framework-system-design.docx"
DIAGRAM = WORKSPACE / "work" / "document-template" / "system-architecture.png"

NAVY = "0B3152"
BLUE = "2878A9"
PALE = "EAF3F8"
INK = "18344B"


def font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\aptos.ttf"),
        Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    bold_candidates = [
        Path(r"C:\Windows\Fonts\aptos-bold.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf"),
    ]
    for candidate in bold_candidates if bold else candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_box(draw, xy, title, body, fill="FFFFFF", outline=NAVY):
    draw.rounded_rectangle(xy, radius=14, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, _ = xy
    draw.text((x1 + 18, y1 + 16), title, font=font(24, True), fill=f"#{INK}")
    y = y1 + 55
    for line in body:
        draw.text((x1 + 18, y), line, font=font(18), fill=f"#{INK}")
        y += 27


def arrow(draw, start, end):
    draw.line([start, end], fill=f"#{BLUE}", width=6)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 11), (x - 18, y + 11)], fill=f"#{BLUE}")


def build_diagram() -> bytes:
    canvas = Image.new("RGB", (1800, 920), "white")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1800, 120), fill=f"#{NAVY}")
    draw.text((60, 30), "ABK-native Vendor-neutral SDD Framework", font=font(34, True), fill="white")
    draw.text((60, 78), "Canonical semantics inside; host syntax and side effects at controlled edges.", font=font(20), fill="#D9E9F3")

    rounded_box(draw, (70, 210, 360, 390), "Entry & Trigger", ["CLI / Agent Host", "User / Hook / Event"])
    rounded_box(draw, (440, 210, 770, 390), "Policy & Resolver", ["Scope + trust validation", "Layer precedence"])
    rounded_box(draw, (850, 210, 1180, 390), "Workflow Engine", ["Typed step algebra", "Gates + bounded loops"])
    rounded_box(draw, (1260, 210, 1710, 390), "ABK Host Adapter", ["Codex empirical host", "Preview = install plan"])
    arrow(draw, (360, 300), (440, 300))
    arrow(draw, (770, 300), (850, 300))
    arrow(draw, (1180, 300), (1260, 300))

    rounded_box(draw, (170, 520, 520, 710), "Global Registry", ["Roles + capabilities", "Adapters + trust policy"], fill=PALE)
    rounded_box(draw, (590, 520, 940, 710), "Project Artifact Graph", ["Spec → plan → evidence", "Owned generated surface"], fill=PALE)
    rounded_box(draw, (1010, 520, 1360, 710), "Session Run Journal", ["Program counter + gates", "Context + evidence"], fill=PALE)
    rounded_box(draw, (1430, 520, 1750, 710), "Local Sandbox", ["Worktree + cache", "Locks + scratch"], fill=PALE)

    draw.line([(345, 520), (605, 390)], fill=f"#{BLUE}", width=4)
    draw.line([(765, 520), (1015, 390)], fill=f"#{BLUE}", width=4)
    draw.line([(1185, 520), (1015, 390)], fill=f"#{BLUE}", width=4)
    draw.line([(1590, 520), (1485, 390)], fill=f"#{BLUE}", width=4)

    draw.rectangle((0, 840, 1800, 920), fill="#D8E8F3")
    draw.text((60, 865), "Invariant: no side effect without pinned inputs, declared capability, durable intent and observable terminal state.", font=font(22, True), fill=f"#{NAVY}")
    canvas.save(DIAGRAM, format="PNG", optimize=True)
    buffer = BytesIO()
    canvas.save(buffer, format="PNG")
    return buffer.getvalue()


def replace_paragraph(paragraph, text: str):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def hard_replace_paragraph(paragraph, text: str):
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if not child.tag.endswith("}pPr"):
            paragraph_element.remove(child)
    if text:
        paragraph.add_run(text)


def replace_exact(document, old: str, new: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old.strip():
            replace_paragraph(paragraph, new)
            return
    raise ValueError(f"Paragraph not found: {old}")


def replace_cell(cell, text: str):
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        replace_paragraph(extra, "")


def fill_table(table, rows):
    if len(rows) != len(table.rows):
        raise ValueError(f"Expected {len(table.rows)} rows, received {len(rows)}")
    for row, values in zip(table.rows, rows):
        if len(values) != len(row.cells):
            raise ValueError("Table column count mismatch")
        for cell, value in zip(row.cells, values):
            replace_cell(cell, value)


def replace_architecture_image(document, image_bytes: bytes):
    shape = document.inline_shapes[0]
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    image_part = document.part.related_parts[blip.embed]
    image_part._blob = image_bytes
    shape._inline.docPr.set(
        "descr",
        "ABK-native vendor-neutral SDD framework architecture: entry and policy boundary feed a typed workflow engine and an ABK host-adapter compiler, supported by Global, Project, Session and Local state layers.",
    )
    shape._inline.docPr.set("title", "Vendor-neutral SDD orchestration architecture")


def update_footer(document):
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                replace_paragraph(paragraph, "Vendor-neutral SDD Framework | System Design RFC")


def main():
    document = Document(REFERENCE)
    image_bytes = build_diagram()
    replace_architecture_image(document, image_bytes)

    replacements = {
        "System Name": "Vendor-neutral SDD Framework",
        "Title of Proposal": "Evidence-based System Design",
        "1.  Abstract": "1.  Összefoglaló",
        "[Summarize the proposed system, the problem it solves, and the intended outcome. Describe the core design at a high level, including the main boundaries, dependencies, and guarantees. Keep this section concise enough that a reviewer can understand the proposal without reading the full document.]": "Az öt vizsgált SDD framework kizárólag mintaforrás. A cél egy ABK-native saját orchestration mag: typed artifact graph, role- és workflow-szerződések, eventek, tartós run journal és kontrollált materialization. A hasznos mechanizmusok saját hook/skill/script/role/workflow/template komponenssé válnak; a közös empirikus host Codex, más host csak későbbi saját host-adapter.",
        "[Describe the workloads and constraints this design must support. Clarify what the proposal does not attempt to solve, the assumptions it relies on, and the most important operational or implementation boundaries.]": "A rendszer ember–agent együttműködésre, sub-agent delegálásra, több sessionön átívelő folytatásra és auditálható projektartefaktumokra optimalizált. Nem helyettesít modellt vagy IDE-t, nem szintetizál jóváhagyást, és nem tekinti a promptot biztonsági sandboxnak. Minden side effect deklarált capabilityhez, rögzített bemenethez és megfigyelhető terminal state-hez kötött. Evidence status: a jelölti viselkedések S1/S2 bizonyítékok; a közös scope-ok, contractok, könyvtárnevek és SLO-k S3 javasolt design.",
        "2.  Goals and Non-Goals": "2.  Célok és nem-célok",
        "3.  Background and Problem Statement": "3.  Háttér és probléma",
        "[Describe the current state, the specific problem, and why the existing approach is no longer sufficient. Include relevant scale, reliability, security, cost, or developer-experience constraints, and explain the impact of leaving the problem unresolved.]": "A jelöltek külön-külön erős mintákat adnak: OpenSpec artifact DAGot; Spec Kit ownershipöt és security utilityket; GSD registryt és producer/checker loopokat; BMAD execution packeteket; Paul PLAN–APPLY–UNIFY ciklust. Upstream runtimeot nem fork-olunk, telepítünk vagy providerként kötünk be: a kiválasztott mechanizmust ABK-native saját komponensként, közös empirikus hostként Codexen mérjük. Paulnál 58 concrete installed reference-nek nincs distributed targetje; tényleges Claude runtime failure nem lett reprodukálva.",
        "[Explain the proposed system boundary and the major responsibilities on each side. Name the primary components, how they interact, and which inputs determine behavior. State the key design principle or invariant that should guide implementation and review.]": "A rendszerhatár belsejében az ABK-native canonical IR, schema, policy, artifact graph és run journal található. A platform-specifikus fájlformátum, hook, parancs és tool invocation a saját host-adapterperem felelőssége. Alapelv: canonical semantics inside, host syntax at the edges; generált fájl csak hash- és ownership-bizonyítékkal írható felül.",
        "4.  Proposed Architecture": "4.  Javasolt architektúra",
        "Figure 1. [Proposed System Architecture].": "1. ábra. Vendorsemleges SDD orchestration architektúra és scope-rétegek.",
        "Core components": "Fő komponensek",
        "5.  Request Lifecycle": "5.  Kérés- és workflow-életciklus",
        "[Describe how a request, event, or job enters the system and identify the required inputs.]": "A User, CLI, native hook vagy agent event létrehoz egy versioned RunEnvelope-ot: event type, scope, actor, correlation ID és kívánt workflow.",
        "[Describe validation, authentication, authorization, and normalization at the system boundary.]": "A boundary guard validálja a schemát, realpath-scope-ot, trust source-ot, engedélyezett capabilityket és az adapter támogatását; bizonytalan scope esetén fail closed.",
        "[Describe which configuration, policy, state, or dependency data is loaded before processing.]": "A javasolt S3 resolver kulcsonként rögzíti a Global/Project/Session/Local forrást és precedence-et; a Global trust/policy kulcsok nem felülírhatók. Ezután pineli az artifact graph verzióit, a role contractot, a package lockot és az ownership manifestet. A logical owner és a fizikai fájlhely eltérhet.",
        "[Describe the primary decision or processing step and the output it produces.]": "A Workflow Engine typed step algebra alapján Sequence, Gate, Branch, bounded Loop, FanOut/FanIn, AgentStep, ToolStep és ArtifactWrite elemeket ütemez.",
        "[Describe which durable state must be written before side effects or downstream execution begin.]": "Side effect előtt atomikusan rögzítendő a program counter, pinned input hashok, idempotency key, approval state, timeout/retry budget és a tervezett state transition.",
        "[Describe downstream calls, retries, timeouts, budget limits, and terminal conditions.]": "Az adapter csak deklarált toolt hívhat. Retry kizárólag besorolt, idempotens hibára engedett; minden loopnak max iteration, deadline és done/blocked/paused/failed terminal state-je van.",
        "[Describe final state updates, the response or output, and the metrics, logs, and traces emitted.]": "A rendszer validálja az outputot, frissíti az artifact graphot és lezárja a run journalt; eventet, metricet, logot, trace-t és evidence summaryt bocsát ki.",
        "6.  API and Data Contracts": "6.  API- és adatszerződések",
        "Primary data contract": "Elsődleges RunEnvelope szerződés",
        "Contract guarantees": "Szerződéses garanciák",
        "[State the ordering, durability, or validation guarantee that must hold before execution.]": "Execution előtt a RunEnvelope, effective config, pinned artifact set és capability decision tartós és schema-valid.",
        "[State how writes, attempts, or events are identified and ordered.]": "Minden write/run/attempt monotonic sequence-et, run ID-t, correlation ID-t és idempotency key-t kap.",
        "[State which versions, identifiers, or inputs must be captured for audit and replay.]": "Audit/replay rögzíti a framework-, schema-, adapter-, role-, workflow-, config- és input-hash verziókat.",
        "[State what this data is and is not a source of truth for.]": "A project artifact graph a delivery state source of truth; a run journal végrehajtási bizonyíték, nem üzleti specifikáció.",
        "The versioned schema or interface definition is published at ": "A versioned schema és adapter contract a projekt `framework/schemas/` és `framework/adapters/` könyvtárában publikált; minden contract release együtt verziózza őket.",
        "[Link to interface or schema] and update it with each contract release.": "",
        "7. Consistency, Idempotency, and Replay": "7. Konzisztencia, idempotencia és replay",
        "[Explain the consistency, idempotency, replay, ordering, or concurrency guarantees required by this design. Distinguish client-facing guarantees from internal analysis or recovery behavior, and identify where duplicate work or partial failure is acceptable.]": "Project-state mutation single-writer lockkal és atomic replace-szel történik. Preview és install ugyanazt az immutable MaterializationPlan objektumot fogyasztja. Replay csak rögzített input- és config-hashokkal engedett; eltérés új run. Párhuzamos agentek külön output ownershipöt kapnak, majd explicit FanIn validálja az összefésülést.",
        "8. Security and Privacy Considerations": "8. Biztonság és adatvédelem",
        "[Describe authentication, authorization, and tenant or data-boundary requirements.]": "A capability policy allow-list alapú; scope, realpath, symlink és worktree ownership minden side effect előtt ellenőrzött.",
        "[Describe data minimization, sensitive payload handling, and logging restrictions.]": "A context manifest csak szükséges artefaktumot tölt; secret, credential, PII és teljes tool output nem kerül promptba vagy logba.",
        "[Describe credential, secret, key, and certificate storage and access requirements.]": "Credentialt a framework nem tárol projektfájlban; platform secret store-ból rövid élettartamú capabilityként kapja.",
        "[Describe safe defaults for administrative, replay, migration, and debugging tools.]": "Network, shell, destructive VCS, migration és external package alapból tiltott; preview, explicit approval, bounded scope és read-back szükséges.",
        "[Describe retention, deletion, residency, privacy, and audit requirements.]": "Run evidence retention policyvezérelt; append-only audit különül a törölhető scratch/cache-től; deletion célpontja explicit és recovery-tervhez kötött.",
        "9.  Operational Readiness": "9.  Üzemeltetési készenlét",
        "10. Alternatives Considered": "10. Vizsgált alternatívák",
        "11. Open Questions": "11. Nyitott döntések",
        "[Open question 1: identify a decision that requires reviewer input before implementation.]": "Az első validáció csak Codex hoston fusson, vagy a core stabilizálása után nyissunk második saját host-adaptert?",
        "[Open question 2: identify an unresolved product, policy, or operational constraint.]": "A human approval ledger személyhez, csapathoz vagy repository policyhoz legyen kötve?",
        "[Open question 3: identify a scale, deployment, or regional design choice.]": "A state maradjon YAML/Markdown, vagy a journal használjon SQLite-ot exportált nézettel?",
        "[Open question 4: identify an ownership, access-control, or tooling decision.]": "Community package trusthez elég a hash + consent, vagy kell szervezeti aláírás?",
        "12. Decision and Next Steps": "12. Döntés és következő lépések",
        "[State the recommended decision and summarize the implementation sequence. Name the first milestone, the validation or dry-run stage, the initial production audience, and the conditions that must be met before broader rollout.]": "Döntés: clean-room, ABK-native vendorsemleges mag, kezdetben Codex hosttal; a jelöltek csak mintaforrások. Upstream tesztpassz nem CHOSEN. CHOSEN csak ugyanazon 10 esetes baseline → eredeti jelölt → ABK-prototípus kapu után; ADOPTED külön approval. Direct reuse esetén az MIT notice jogi kötelezettség, a repository/commit/file provenance külön belső auditpolicy. Író pilot csak parity-, security-, recovery- és emberi approval-bizonyíték után indulhat.",
    }

    for old, new in replacements.items():
        replace_exact(document, old, new)

    for paragraph in document.paragraphs:
        if "[Link to interface or schema]" in paragraph.text:
            hard_replace_paragraph(paragraph, "")
        if paragraph.text.strip() == "4.  Javasolt architektúra":
            paragraph.paragraph_format.page_break_before = True
        if paragraph.text.strip() == "Community package trusthez elég a hash + consent, vagy kell szervezeti aláírás?":
            paragraph.paragraph_format.space_before = Pt(2)

    fill_table(document.tables[0], [["STATUS\nProposed", "", "OWNER\nFramework Architecture", "", "LAST UPDATED\nAugust 02, 2026"]])
    fill_table(document.tables[1], [
        ["Szerzők", "OpenAI Codex — source-backed multi-agent research"],
        ["Reviewerek", "Framework owner; security; developer experience; platform adapters"],
        ["Kapcsolódó anyagok", "Öt jelölti mintaforrás-dosszié (S1/S2 evidence) és reusable-pattern katalógus (S3 ABK-native mapping)"],
        ["Scope", "Javasolt S3 ABK-native Global–Project–Session–Local core, saját Codex host-adapter, state, security és operations"],
    ])
    fill_table(document.tables[2], [
        ["Célok", "Nem-célok"],
        ["Deterministic artifact graph és resumable workflow állapot", "Saját foundation model vagy IDE készítése"],
        ["Vendorsemleges canonical IR és tesztelt adapterek", "A promptot biztonsági sandboxként kezelni"],
        ["Auditálható role, capability, event és approval contract", "Automatikus jóváhagyás, push, merge vagy production módosítás"],
        ["Safe preview/install/update/recovery/uninstall lifecycle", "Harmadik fél teljes frameworkjének változatlan átvétele"],
    ])
    fill_table(document.tables[3], [
        ["Komponens", "Felelősség", "Elsődleges tárolás", "Hibaviselkedés"],
        ["Canonical Registry & IR", "Schema, role, workflow, event, capability és package verziók", "Global registry + project lock", "Ismeretlen verzió vagy reference: fail closed"],
        ["Policy & Resolver", "Layer precedence, trust, scope, consent és effective config", "Policy store + project manifest", "Bizonytalan scope/permission: execution tiltás"],
        ["Workflow & Event Engine", "Typed step algebra, program counter, gate, bounded loop, fan-out/fan-in", "Session run journal", "Atomikus checkpoint; explicit paused/failed state"],
        ["ABK Host Adapter Compiler", "Preview/materialize/update/remove; Codex first, később saját host-adapter", "Staging + ownership manifest", "Unowned overwrite/symlink/drift: abort és rollback"],
        ["Artifact, Evidence & Observability", "Project DAG, evidence, metrics, logs, traces és recovery", "Project state + append-only audit", "State-write hiba előtt nincs downstream side effect"],
    ])
    fill_table(document.tables[4], [
        ["Mező", "Típus", "Kötelező", "Leírás"],
        ["schema_version", "SemVer", "Igen", "RunEnvelope és transition contract verzió; ismeretlen major elutasítandó."],
        ["run_id / correlation_id", "UUID", "Igen", "Egy futás és a kapcsolódó eventlánc stabil azonosítói."],
        ["scope", "Enum", "Igen", "Global, Project, Session vagy Local; realpath-bound ownershipdel."],
        ["event", "Object", "Igen", "Type, source, actor, timestamp, preconditions és requested workflow."],
        ["pinned_inputs", "Hash map", "Igen", "Artifact, config, role, workflow, adapter és package verziók/hashok."],
        ["capabilities", "Array", "Igen", "Engedélyezett tools, write scope, network, secrets, timeout és retry class."],
        ["transition", "Object", "Igen", "From/to state, idempotency key, attempt, evidence és terminal reason."],
    ])
    fill_table(document.tables[5], [
        ["Scenario", "Várt viselkedés", "Indoklás"],
        ["Duplikált event vagy replay", "Az idempotency key ugyanazt a terminal resultot adja; nincs új side effect.", "A program counter és output ownership tartós."],
        ["Kötelező state write vagy dependency hibázik", "Fail closed; retryability besorolt; downstream végrehajtás nem indul.", "A durable intent megelőzi a side effectet."],
        ["Timeout vagy részleges adapterhiba", "Bounded retry után failed/paused; journal recovery packetet tartalmaz.", "A részállapot látható és rekonstruálható."],
        ["Config változik futás közben", "A futás a pinned effective configot használja; új verzió új run.", "Nincs mid-run semantic drift."],
    ])
    fill_table(document.tables[6], [
        ["Signal", "SLO vagy alert", "Owner", "Launch gate"],
        ["Workflow terminal completion", ">=99% kontrollált tesztcorpuson; 0 néma success", "Core runtime", "Kötelező"],
        ["Checkpoint/materialization latency", "P95 cél és workflow-class timeout dokumentált", "Runtime + adapters", "Ajánlott"],
        ["Retry/fallback/error rate", "Retry class és terminal reason szerint; fallback nem lehet néma", "Ops", "Kötelező"],
        ["Generated/config drift", "0 ownership-, hash-, registry- és adapter-parity eltérés", "ABK host adapters", "Kötelező"],
        ["Recovery/replay integrity", "Crash-injection és restore teszt minden release-ben", "Core + security", "Kötelező"],
        ["Rollout constraint: read-only dry-run → disposable write pilot → one-project canary → bounded organization rollout; rollback and uninstall rehearsal precedes each promotion.", "", "", ""],
    ])
    fill_table(document.tables[7], [
        ["Mintaforrás", "Átvett gear", "ABK-native cél"],
        ["OpenSpec", "Artifact DAG és root provenance minta.", "Saját artifact graph; nincs upstream fork/runtime."],
        ["GitHub Spec Kit", "Ownership-, adapter- és security minták.", "Saját hook/skill/script és Codex host-adapter."],
        ["GSD és BMAD", "Agent-, role-, loop-, registry- és workflow minták.", "Saját role/workflow/template; nincs teljes port."],
        ["Paul", "PLAN–APPLY–UNIFY ciklus.", "Saját workflow; 58 target nélküli ref és CARL-hiány nem öröklődik."],
    ])
    fill_table(document.tables[8], [
        ["Mérföldkő", "Deliverable", "Kilépési feltétel"],
        ["M1", "Canonical schema, artifact graph, run journal, read-only CLI", "Schema, DAG readiness és replay contract tesztek zöldek"],
        ["M2", "Saját Codex host-adapter + pure MaterializationPlan dry-run", "Preview/install parity, ownership, path/symlink és rollback tesztek zöldek"],
        ["M3", "Role/sub-agent engine, gates, bounded loops, evidence", "Disposable end-to-end pilot; 0 néma fallback; recovery packet rekonstruálható"],
        ["M4", "Második saját host-adapter, package catalog és operations", "Canary SLO, uninstall/recovery rehearsal és security approval teljes"],
    ])

    update_footer(document)
    document.core_properties.title = "Vendor-neutral SDD Framework — Evidence-based System Design"
    document.core_properties.subject = "Global–Project–Session–Local agentic orchestration architecture"
    document.core_properties.author = "OpenAI Codex"
    document.core_properties.keywords = "SDD, agent, workflow, Codex, Claude, system design"
    document.core_properties.comments = "Generated from the retained System Design reference template."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
